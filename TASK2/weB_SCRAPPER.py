import os
from time import sleep, time
import requests
from bs4 import BeautifulSoup
from docx import Document
import tkinter as tk
from tkinter import ttk
from tkinter import messagebox
from io import BytesIO
from PIL import Image

def get_page(url):
    headers = {
        'User-Agent': '[sujal]SkillSync internship and personel study',
    }

    response = requests.get(url, headers=headers)
    
    if response.status_code == 200:
        return response.text
    else:
        print(f"Failed to retrieve the page. Status Code: {response.status_code}")
        return None

def download_image(img_url, output_dir):
    try:
        response = requests.get(img_url)
        img_data = BytesIO(response.content)
        img = Image.open(img_data)
        img_name = os.path.join(output_dir, f"image_{len(os.listdir(output_dir)) + 1}.png")
        img.save(img_name)
        return img_name
    except Exception as e:
        print(f"Failed to download image: {e}")
        return None

def extract_content(html_content, output_dir):
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Create a Word document
    doc = Document()

    # Extract text and add it to the Word document
    text = soup.get_text(separator=' ', strip=True)
    doc.add_paragraph(text)

    # Extract and download images
    for img_tag in soup.find_all('img'):
        img_url = img_tag.get('src')
        img_path = download_image(img_url, output_dir)
        
        if img_path:
            # Add image to the Word document
            doc.add_picture(img_path)
            os.remove(img_path)  # Remove the downloaded image file

    return doc

def scrape_and_save(base_url, start_page, num_pages, rate_limit, output_dir):
    for i in range(1, num_pages + 1):
        url = f"{base_url}{start_page}?page={i}"
        
        # Respectful crawling: Adhere to rate limits
        sleep(rate_limit)
        
        html_content = get_page(url)
        
        if html_content:
            doc = extract_content(html_content, output_dir)
            doc.save(os.path.join(output_dir, f"page_{i}_content.docx"))

def on_submit():
    base_url = entry.get()
    num_pages = int(num_pages_var.get())
    rate_limit = int(rate_limit_var.get())
    output_dir = "website_content"
    
    if base_url:
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        scrape_and_save(base_url, '/start', num_pages, rate_limit, output_dir)
        messagebox.showinfo("Scraping Complete", "Content has been retrieved and saved to 'website_content' folder.")
    else:
        messagebox.showerror("Error", "Please enter a valid URL.")

# Create the main window
root = tk.Tk()
root.title("Web Content Scraper UI")

# Create and place widgets in the window
label = tk.Label(root, text="Enter Website URL:")
label.pack(pady=10)

entry = tk.Entry(root, width=40)
entry.pack(pady=10)

label_num_pages = tk.Label(root, text="Number of Pages:")
label_num_pages.pack(pady=5)

num_pages_var = tk.StringVar()
num_pages_var.set("1")
entry_num_pages = tk.Entry(root, textvariable=num_pages_var)
entry_num_pages.pack(pady=10)

label_rate_limit = tk.Label(root, text="Rate Limit (seconds):")
label_rate_limit.pack(pady=5)

rate_limit_var = tk.StringVar()
rate_limit_var.set("2")
entry_rate_limit = tk.Entry(root, textvariable=rate_limit_var)
entry_rate_limit.pack(pady=10)

submit_button = tk.Button(root, text="Submit", command=on_submit)
submit_button.pack(pady=20)

# Start the Tkinter event loop
root.mainloop()
